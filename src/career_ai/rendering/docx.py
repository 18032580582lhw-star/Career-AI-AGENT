"""DOCX renderer for accepted structured resume documents."""

from __future__ import annotations

from typing import TYPE_CHECKING, assert_never, final

from docx import Document

from career_ai.rendering.artifacts import output_artifact
from career_ai.rendering.models import RendererOutcome, RendererSuccess
from career_ai.tailoring.document_contracts import (
    AcceptedResumeDocument,
    EducationEntry,
    ExperienceEntry,
    ProjectEntry,
    ResumeSection,
)
from career_ai.tailoring.manifest_contracts import RenderBackend

if TYPE_CHECKING:
    from pathlib import Path

    from docx.document import Document as DocxDocument

_DOCX_NAME = "resume.docx"
_DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


@final
class DocxResumeRenderer:
    """Render accepted resume documents to WordprocessingML DOCX."""

    @property
    def backend(self) -> RenderBackend:
        """Return the registry backend implemented by this adapter."""
        return RenderBackend.DOCX

    def render(
        self,
        document: AcceptedResumeDocument,
        output_directory: Path,
    ) -> RendererOutcome:
        """Write one DOCX artifact for an accepted resume document."""
        output_directory.mkdir(parents=True, exist_ok=True)
        output_path = output_directory / _DOCX_NAME
        docx: DocxDocument = Document()
        _add_identity(docx, document)
        for section in document.section_order:
            _add_section(docx, document, section)
        docx.save(str(output_path))
        return RendererSuccess(
            backend=self.backend,
            artifacts=(
                output_artifact(
                    output_path,
                    relative_path=_DOCX_NAME,
                    media_type=_DOCX_MEDIA_TYPE,
                ),
            ),
        )


def _add_identity(docx: DocxDocument, document: AcceptedResumeDocument) -> None:
    _ = docx.add_heading(document.identity.name, level=0)
    if document.identity.headline is not None:
        _ = docx.add_paragraph(document.identity.headline)
    for contact in document.identity.contact_lines:
        _ = docx.add_paragraph(contact)


def _add_section(
    docx: DocxDocument,
    document: AcceptedResumeDocument,
    section: ResumeSection,
) -> None:
    match section:
        case ResumeSection.SUMMARY:
            _add_bulleted_section(
                docx,
                "Professional Summary",
                tuple(bullet.text for bullet in document.professional_summary),
            )
        case ResumeSection.SKILLS:
            _add_bulleted_section(
                docx,
                "Skills",
                tuple(skill.text for skill in document.skills),
            )
        case ResumeSection.EXPERIENCE:
            _add_experience_section(docx, document)
        case ResumeSection.PROJECTS:
            _add_project_section(docx, document)
        case ResumeSection.EDUCATION:
            _add_education_section(docx, document)
        case ResumeSection.LINKS:
            _add_link_section(docx, document)
        case _:
            assert_never(section)


def _add_bulleted_section(
    docx: DocxDocument,
    title: str,
    items: tuple[str, ...],
) -> None:
    _ = docx.add_heading(title, level=1)
    for item in items:
        _ = docx.add_paragraph(item, style="List Bullet")


def _add_experience_section(
    docx: DocxDocument,
    document: AcceptedResumeDocument,
) -> None:
    _ = docx.add_heading("Experience", level=1)
    for entry in document.experience:
        _add_experience(docx, entry)


def _add_experience(docx: DocxDocument, entry: ExperienceEntry) -> None:
    _ = docx.add_heading(f"{entry.title}, {entry.organization}", level=2)
    details = (
        entry.date_range
        if entry.location is None
        else f"{entry.date_range} | {entry.location}"
    )
    _ = docx.add_paragraph(details)
    for bullet in entry.bullets:
        _ = docx.add_paragraph(bullet.text, style="List Bullet")


def _add_project_section(
    docx: DocxDocument,
    document: AcceptedResumeDocument,
) -> None:
    _ = docx.add_heading("Projects", level=1)
    for entry in document.projects:
        _add_project(docx, entry)


def _add_project(docx: DocxDocument, entry: ProjectEntry) -> None:
    _ = docx.add_heading(entry.name, level=2)
    if entry.subtitle is not None:
        _ = docx.add_paragraph(entry.subtitle)
    for bullet in entry.bullets:
        _ = docx.add_paragraph(bullet.text, style="List Bullet")


def _add_education_section(
    docx: DocxDocument,
    document: AcceptedResumeDocument,
) -> None:
    _ = docx.add_heading("Education", level=1)
    for entry in document.education:
        _add_education(docx, entry)


def _add_education(docx: DocxDocument, entry: EducationEntry) -> None:
    _ = docx.add_heading(entry.institution, level=2)
    credential = entry.credential
    if entry.date_range is not None:
        credential = f"{credential} | {entry.date_range}"
    _ = docx.add_paragraph(credential)
    for detail in entry.details:
        _ = docx.add_paragraph(detail.text, style="List Bullet")


def _add_link_section(
    docx: DocxDocument,
    document: AcceptedResumeDocument,
) -> None:
    _ = docx.add_heading("Links", level=1)
    for link in document.links:
        _ = docx.add_paragraph(f"{link.label}: {link.url}")
