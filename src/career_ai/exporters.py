from pathlib import Path

from docx import Document

from career_ai.models import CareerFitReport


def build_resume_docx(report: CareerFitReport, output_path: Path) -> Path:
    """Build a tailored resume document and return its path."""
    document = Document()
    _ = document.add_heading("Tailored Resume", level=1)
    _ = document.add_paragraph(f"Target role: {report.jd_analysis.role_title}")
    _ = document.add_heading("Match Summary", level=2)
    _ = document.add_paragraph(f"Match score: {report.match.score}")
    _ = document.add_paragraph(report.match.summary)
    _ = document.add_heading("Recommended Bullets", level=2)
    for suggestion in report.bullet_suggestions:
        _ = document.add_paragraph(suggestion.improved, style="List Bullet")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def build_cover_letter_docx(report: CareerFitReport, output_path: Path) -> Path:
    """Build a cover letter document and return its path."""
    document = Document()
    for paragraph in report.cover_letter_draft.split("\n\n"):
        _ = document.add_paragraph(paragraph)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
