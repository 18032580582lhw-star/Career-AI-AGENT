import hashlib
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from career_ai.workspace import create_workspace
from career_ai.workspace.ingestion import (
    IngestionError,
    IngestionErrorCode,
    ingest_jd_text,
    ingest_jd_url,
    ingest_latex_template,
    ingest_resume_file,
)


def test_txt_resume_is_content_addressed_and_duplicate_is_identity(tmp_path: Path) -> None:
    _ = create_workspace(tmp_path)
    resume = tmp_path / "resume.txt"
    original = b"Python engineer\nBuilt safe systems.\n"
    _ = resume.write_bytes(original)

    first = ingest_resume_file(tmp_path, resume)
    second = ingest_resume_file(tmp_path, resume)

    assert first == second
    assert first.sha256 == hashlib.sha256(original).hexdigest()
    assert (tmp_path / first.content_path).read_bytes() == original
    assert (tmp_path / first.extracted_text_path).read_text(encoding="utf-8") == original.decode()


def test_docx_resume_extracts_paragraph_text(tmp_path: Path) -> None:
    _ = create_workspace(tmp_path)
    resume = tmp_path / "resume.docx"
    document = Document()
    _ = document.add_paragraph("Data Engineer")
    _ = document.add_paragraph("Built ETL pipelines")
    document.save(str(resume))

    artifact = ingest_resume_file(tmp_path, resume)

    extracted = (tmp_path / artifact.extracted_text_path).read_text(encoding="utf-8")
    assert artifact.media_type.endswith("wordprocessingml.document")
    assert extracted == "Data Engineer\nBuilt ETL pipelines"


def test_malformed_docx_reports_typed_source_read_failure(tmp_path: Path) -> None:
    _ = create_workspace(tmp_path)
    resume = tmp_path / "malformed.docx"
    _ = resume.write_bytes(b"not-a-zip-package")

    with pytest.raises(IngestionError) as captured:
        _ = ingest_resume_file(tmp_path, resume)

    assert captured.value.code is IngestionErrorCode.SOURCE_READ_FAILED


def test_blank_pdf_reports_scanned_pdf_without_ocr(tmp_path: Path) -> None:
    _ = create_workspace(tmp_path)
    resume = tmp_path / "scan.pdf"
    writer = PdfWriter()
    _ = writer.add_blank_page(width=100, height=100)
    with resume.open("wb") as stream:
        _ = writer.write(stream)

    with pytest.raises(IngestionError) as captured:
        _ = ingest_resume_file(tmp_path, resume)

    assert captured.value.code is IngestionErrorCode.SCANNED_PDF


def test_encrypted_pdf_reports_explicit_error(tmp_path: Path) -> None:
    _ = create_workspace(tmp_path)
    resume = tmp_path / "encrypted.pdf"
    writer = PdfWriter()
    _ = writer.add_blank_page(width=100, height=100)
    writer.encrypt("secret")
    with resume.open("wb") as stream:
        _ = writer.write(stream)

    with pytest.raises(IngestionError) as captured:
        _ = ingest_resume_file(tmp_path, resume)

    assert captured.value.code is IngestionErrorCode.ENCRYPTED_PDF


def test_latex_template_is_utf8_structural_and_never_modified(tmp_path: Path) -> None:
    _ = create_workspace(tmp_path)
    template = tmp_path / "resume.tex"
    original = (
        "\\documentclass{article}\n\\begin{document}\n"
        "\\section{Experience}原始内容\n\\end{document}\n"
    ).encode()
    _ = template.write_bytes(original)

    artifact = ingest_latex_template(tmp_path, template)

    assert template.read_bytes() == original
    assert (tmp_path / artifact.content_path).read_bytes() == original
    assert artifact.kind == "latex_template"


def test_latex_template_rejects_invalid_utf8_and_missing_boundary(tmp_path: Path) -> None:
    _ = create_workspace(tmp_path)
    invalid_encoding = tmp_path / "invalid.tex"
    _ = invalid_encoding.write_bytes(b"\\documentclass{article}\xff")
    with pytest.raises(IngestionError) as encoding_error:
        _ = ingest_latex_template(tmp_path, invalid_encoding)
    assert encoding_error.value.code is IngestionErrorCode.INVALID_ENCODING

    missing_boundary = tmp_path / "missing.tex"
    _ = missing_boundary.write_text("\\documentclass{article}", encoding="utf-8")
    with pytest.raises(IngestionError) as structure_error:
        _ = ingest_latex_template(tmp_path, missing_boundary)
    assert structure_error.value.code is IngestionErrorCode.INVALID_LATEX


def test_jd_text_and_hardened_url_are_stored_as_jd_sources(tmp_path: Path) -> None:
    _ = create_workspace(tmp_path)

    text_artifact = ingest_jd_text(tmp_path, "Required: Python and SQL")
    url_artifact = ingest_jd_url(
        tmp_path,
        "data:text/html,<main><h1>Engineer</h1><p>Required: Python</p></main>",
    )

    assert text_artifact.kind == "job_description"
    assert text_artifact.origin == "text"
    assert url_artifact.origin == "url"
    assert "Engineer" in (tmp_path / url_artifact.extracted_text_path).read_text(encoding="utf-8")


def test_jd_url_preserves_hardened_fetch_rejection(tmp_path: Path) -> None:
    _ = create_workspace(tmp_path)

    with pytest.raises(IngestionError) as captured:
        _ = ingest_jd_url(tmp_path, "http://127.0.0.1/private")

    assert captured.value.code is IngestionErrorCode.JD_FETCH_FAILED
    assert "blocked_host" in captured.value.reason


def test_unsupported_suffix_cannot_influence_workspace_output_path(tmp_path: Path) -> None:
    _ = create_workspace(tmp_path)
    source = tmp_path / "..evil.exe"
    _ = source.write_bytes(b"not a resume")

    with pytest.raises(IngestionError) as captured:
        _ = ingest_resume_file(tmp_path, source)

    assert captured.value.code is IngestionErrorCode.UNSUPPORTED_MEDIA_TYPE
    assert not list((tmp_path / ".career_ai" / "sources").rglob("*.exe"))
